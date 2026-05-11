"""Generates FSD.docx and TSD.docx for the CRQ Approval Automation System."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour palette ────────────────────────────────────────────────────────────
BRAND_BLUE   = RGBColor(0x1F, 0x49, 0x7D)   # dark blue  – headings / header row
ACCENT_BLUE  = RGBColor(0x2E, 0x75, 0xB6)   # mid blue   – sub-headings
LIGHT_BLUE   = RGBColor(0xD6, 0xE4, 0xF0)   # pale blue  – table header fill
ALT_ROW      = RGBColor(0xF2, 0xF7, 0xFD)   # very light – alternate table rows
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
BLACK        = RGBColor(0x00, 0x00, 0x00)
GRAY         = RGBColor(0x59, 0x59, 0x59)
DARK_GRAY    = RGBColor(0x26, 0x26, 0x26)

# ─────────────────────────────────────────────────────────────────────────────
# Helper utilities
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="BFBFBF"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tblBorders.append(b)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a styled table with blue header row and alternating row colours."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, BRAND_BLUE)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = ALT_ROW if r_idx % 2 == 1 else WHITE
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_bg(cell, bg)
            para = cell.paragraphs[0]
            run = para.add_run(str(cell_text) if cell_text is not None else "")
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            run.font.color.rgb = DARK_GRAY

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    return table


def heading1(doc, text):
    p = doc.add_paragraph()
    p.style = "Heading 1"
    run = p.add_run(text)
    run.font.color.rgb = BRAND_BLUE
    run.font.size = Pt(16)
    run.font.name = "Calibri"
    run.bold = True
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    p.style = "Heading 2"
    run = p.add_run(text)
    run.font.color.rgb = ACCENT_BLUE
    run.font.size = Pt(13)
    run.font.name = "Calibri"
    run.bold = True
    return p


def heading3(doc, text):
    p = doc.add_paragraph()
    p.style = "Heading 3"
    run = p.add_run(text)
    run.font.color.rgb = ACCENT_BLUE
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.bold = True
    return p


def body(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Calibri"
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color if color else DARK_GRAY
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK_GRAY
    return p


def info_box(doc, label, value):
    """Single-row 2-column info line (used in cover / meta section)."""
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}:  ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.name = "Calibri"
    r1.font.color.rgb = BRAND_BLUE
    r2 = p.add_run(value)
    r2.font.size = Pt(10.5)
    r2.font.name = "Calibri"
    r2.font.color.rgb = DARK_GRAY
    return p


def cover_page(doc, title, subtitle, version, date, status):
    """Styled cover page."""
    # Top colour band – a 1-cell table acting as a banner
    banner = doc.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = banner.cell(0, 0)
    set_cell_bg(cell, BRAND_BLUE)
    cell.width = Inches(6.5)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run(title)
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = WHITE
    r.font.name = "Calibri"
    para.paragraph_format.space_before = Pt(20)
    para.paragraph_format.space_after = Pt(20)

    doc.add_paragraph()

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p.add_run(subtitle)
    r2.font.size = Pt(16)
    r2.font.color.rgb = ACCENT_BLUE
    r2.font.name = "Calibri"
    r2.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Meta table
    meta = doc.add_table(rows=3, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(meta)
    meta_data = [("Version", version), ("Date", date), ("Status", status)]
    for i, (lbl, val) in enumerate(meta_data):
        lc = meta.cell(i, 0)
        vc = meta.cell(i, 1)
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lr = lp.add_run(lbl + "  ")
        lr.bold = True
        lr.font.size = Pt(11)
        lr.font.color.rgb = BRAND_BLUE
        lr.font.name = "Calibri"
        vp = vc.paragraphs[0]
        vr = vp.add_run(val)
        vr.font.size = Pt(11)
        vr.font.color.rgb = DARK_GRAY
        vr.font.name = "Calibri"

    doc.add_page_break()


def divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E75B6")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def code_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    shading = OxmlElement("w:rPr")
    return p


# ─────────────────────────────────────────────────────────────────────────────
# FSD
# ─────────────────────────────────────────────────────────────────────────────

def build_fsd():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Cover
    cover_page(
        doc,
        "Functional Specification Document",
        "CRQ Approval Automation System",
        "1.0", "2026-05-11", "Draft"
    )

    # ── 1. Purpose ────────────────────────────────────────────────────────────
    heading1(doc, "1.  Purpose")
    divider(doc)
    body(doc,
        "This document describes the functional requirements of the CRQ Approval Automation System. "
        "It is intended for business stakeholders, operations teams, project managers, and quality "
        "assurance personnel who need to understand what the system does, how users interact with it, "
        "and what rules govern its behaviour — without needing to understand the underlying technology.")

    # ── 2. Background ─────────────────────────────────────────────────────────
    heading1(doc, "2.  Background and Business Context")
    divider(doc)
    body(doc,
        "Change Request (CRQ) records are managed in the Remedy IT Service Management (ITSM) platform. "
        "A master list of active CRQs is maintained by the Change Management team in a Microsoft Excel "
        "spreadsheet stored on OneDrive.")
    doc.add_paragraph()
    body(doc,
        "Previously, the operations team had to manually log into Remedy each afternoon, check the "
        "status of every CRQ, identify which ones had been approved, and send notification emails by "
        "hand. This process was slow, inconsistent, and prone to missed notifications.")
    doc.add_paragraph()
    body(doc, "This system eliminates that manual effort by:")
    bullet(doc, "Automatically reading the CRQ list from OneDrive each day at 5:30 PM")
    bullet(doc, "Checking each CRQ's current status against the Remedy API")
    bullet(doc, "Sending a single consolidated notification email for all newly approved CRQs")
    bullet(doc, "Providing a web dashboard for monitoring run history and triggering on-demand runs")

    # ── 3. Stakeholders ───────────────────────────────────────────────────────
    heading1(doc, "3.  Stakeholders")
    divider(doc)
    add_styled_table(doc,
        ["Role", "Responsibility"],
        [
            ["Operations Team",    "Receives approval notification emails; may trigger ad-hoc runs via the web UI"],
            ["Change Manager",     "Maintains the OneDrive Excel file containing the CRQ list"],
            ["System Administrator", "Configures credentials, scheduler settings, and deploys the application"],
        ],
        col_widths=[2.0, 4.5]
    )

    # ── 4. Functional Requirements ────────────────────────────────────────────
    heading1(doc, "4.  Functional Requirements")
    divider(doc)

    heading2(doc, "4.1  Scheduled Daily CRQ Processing")
    add_styled_table(doc,
        ["Req. ID", "Description"],
        [
            ["FR-01", "The system shall automatically trigger a CRQ processing run every day at 5:30 PM."],
            ["FR-02", "On each scheduled run, the system shall download the current CRQ Excel file from the configured OneDrive location."],
            ["FR-03", "The system shall parse all CRQ rows from the Excel file (columns: CRQ Number, Title, Assignee, Description, Last Updated). The first row is a header and is skipped. Rows with a blank CRQ number are also skipped."],
            ["FR-04", "For each CRQ, the system shall query the Remedy REST API to retrieve the current status."],
            ["FR-05", "A CRQ is considered approved if its Remedy status matches the configured approval status value (default: \"Request in Change\")."],
            ["FR-06", "The system shall send a single consolidated email listing all approved CRQs. If no CRQs are approved, no email is sent."],
            ["FR-07", "After sending the email, the system shall record the email-sent flag and timestamp against each approved CRQ."],
            ["FR-08", "After each run — whether successful or not — the system shall create a processing log entry recording the outcome."],
        ],
        col_widths=[1.0, 5.5]
    )

    heading2(doc, "4.2  Ad-Hoc CRQ Processing")
    add_styled_table(doc,
        ["Req. ID", "Description"],
        [
            ["FR-09", "An authorised user shall be able to trigger a CRQ processing run manually from the web UI at any time."],
            ["FR-10", "The ad-hoc run form shall require: (a) Name / User ID of the person triggering the run, (b) From Date-Time, and (c) To Date-Time."],
            ["FR-11", "An ad-hoc run shall only process CRQs whose Last Updated timestamp in the Excel file falls within the specified From–To date-time range (inclusive)."],
            ["FR-12", "The ad-hoc run shall follow the same logic as the scheduled run (status check, email, logging) but is recorded with batch type ADHOC and the triggering user's name."],
            ["FR-13", "The UI shall validate: Name is not blank; both dates are provided; From date-time is earlier than To date-time. Submission is blocked if validation fails."],
            ["FR-14", "After the run completes, the UI shall display a result summary: run time, type, total CRQs read, approved count, emails sent, and overall status."],
        ],
        col_widths=[1.0, 5.5]
    )

    heading2(doc, "4.3  Email Notification")
    add_styled_table(doc,
        ["Req. ID", "Description"],
        [
            ["FR-15", "The approval email shall be a consolidated notification listing all approved CRQs from the current run."],
            ["FR-16", "The recipient email address(es) shall be configurable by the administrator."],
            ["FR-17", "The sender address and email subject line shall both be configurable by the administrator."],
        ],
        col_widths=[1.0, 5.5]
    )

    heading2(doc, "4.4  Dashboard")
    add_styled_table(doc,
        ["Req. ID", "Description"],
        [
            ["FR-18", "The dashboard shall display today's summary statistics: Total CRQs processed, Approved, Emails Sent, and Pending (not approved)."],
            ["FR-19", "The dashboard shall display the timestamp of the last scheduled run and the last ad-hoc run."],
            ["FR-20", "The dashboard shall show a table of all CRQs processed today with their approval and email status."],
            ["FR-21", "The dashboard shall show the 10 most recent processing log entries."],
            ["FR-22", "A Refresh button shall reload all dashboard data without navigating away."],
        ],
        col_widths=[1.0, 5.5]
    )

    heading2(doc, "4.5  CRQ List")
    add_styled_table(doc,
        ["Req. ID", "Description"],
        [
            ["FR-23", "A dedicated screen shall show all CRQ records across all runs."],
            ["FR-24", "The list shall support free-text search across CRQ number, title, and assignee (case-insensitive)."],
            ["FR-25", "The list shall support status filters: ALL, APPROVED, PENDING, EMAIL_SENT."],
            ["FR-26", "The list shall display a record count showing how many records match the current filter out of the total."],
        ],
        col_widths=[1.0, 5.5]
    )

    heading2(doc, "4.6  Processing Logs")
    add_styled_table(doc,
        ["Req. ID", "Description"],
        [
            ["FR-27", "A dedicated screen shall show all processing log entries, most recent first."],
            ["FR-28", "Each log entry shall show: run timestamp, batch type, total CRQs read, approved count, emails sent, status, triggered-by, and any error message."],
        ],
        col_widths=[1.0, 5.5]
    )

    # ── 5. Use Cases ──────────────────────────────────────────────────────────
    heading1(doc, "5.  Use Cases")
    divider(doc)

    for uc in [
        {
            "id": "UC-01", "name": "Scheduled Daily Run",
            "actor": "System (automatic scheduler)",
            "trigger": "Clock reaches 5:30 PM",
            "pre": "Application is running; OneDrive file is accessible; Remedy API is reachable",
            "flow": [
                "Scheduler fires and calls the processing service.",
                "Excel file is downloaded from OneDrive.",
                "All CRQ rows are parsed from the file.",
                "For each CRQ, the Remedy API is queried for the current status.",
                "Each CRQ record is saved to the database.",
                "All approved CRQs are collected; one consolidated email is sent.",
                "Email-sent flag and timestamp are recorded on each approved CRQ.",
                "A processing log entry is created (SUCCESS / PARTIAL / FAILED).",
            ],
            "alt": [
                "Remedy API error for a single CRQ: error is logged per CRQ; remaining CRQs are processed; log is recorded as PARTIAL.",
                "OneDrive unavailable: exception caught; log recorded as FAILED; no email sent.",
            ],
        },
        {
            "id": "UC-02", "name": "Ad-Hoc Run from UI",
            "actor": "Operations Team member",
            "trigger": "User submits the Ad-Hoc Approval form",
            "pre": "User has entered a valid name and a valid From–To date-time range",
            "flow": [
                "User completes and submits the Ad-Hoc Approval form.",
                "UI validates inputs (name not blank; valid date range).",
                "Request is sent to the backend with triggeredBy, fromDateTime, toDateTime.",
                "Backend filters Excel rows to those whose Last Updated falls within the range.",
                "Same processing pipeline runs on the filtered rows (status check, email, log).",
                "Result summary is returned and displayed in the UI.",
            ],
            "alt": [
                "No CRQs match the date range: run completes with zero results; no email sent; log recorded as SUCCESS with 0 counts.",
            ],
        },
        {
            "id": "UC-03", "name": "Monitor Dashboard",
            "actor": "Operations Team member",
            "trigger": "User opens the application or clicks Refresh",
            "pre": "Application is running",
            "flow": [
                "UI sends a dashboard data request to the backend.",
                "Backend returns today's counts, last run timestamps, recent CRQs, and recent logs.",
                "Dashboard renders all data.",
            ],
            "alt": [],
        },
        {
            "id": "UC-04", "name": "Search and Filter CRQ Records",
            "actor": "Operations Team member",
            "trigger": "User navigates to the CRQ List screen",
            "pre": "Application is running",
            "flow": [
                "All CRQ records are loaded from the backend when the screen opens.",
                "User types in the search box — list filters by CRQ number, title, or assignee.",
                "User clicks a filter button (APPROVED / PENDING / EMAIL_SENT) — list narrows accordingly.",
                "Record count updates to show filtered vs total.",
            ],
            "alt": [],
        },
    ]:
        heading2(doc, f"{uc['id']} — {uc['name']}")
        info_box(doc, "Actor",     uc["actor"])
        info_box(doc, "Trigger",   uc["trigger"])
        info_box(doc, "Pre-condition", uc["pre"])
        body(doc, "Main Flow:", bold=True)
        for i, step in enumerate(uc["flow"], 1):
            bullet(doc, f"{i}.  {step}")
        if uc["alt"]:
            body(doc, "Alternate / Exception Flows:", bold=True)
            for a in uc["alt"]:
                bullet(doc, a)
        doc.add_paragraph()

    # ── 6. Business Rules ─────────────────────────────────────────────────────
    heading1(doc, "6.  Business Rules")
    divider(doc)
    add_styled_table(doc,
        ["Rule ID", "Description"],
        [
            ["BR-01", "A CRQ is approved if and only if its Remedy status equals the configured approval status value."],
            ["BR-02", "Only one consolidated email is sent per batch run; individual per-CRQ emails are not used."],
            ["BR-03", "Ad-hoc runs filter on the CRQ's Last Updated timestamp from the Excel file — not on the current time or processed time."],
            ["BR-04", "The first row of the Excel file is always treated as a header and is never processed as a CRQ."],
            ["BR-05", "Rows with a blank CRQ number in the Excel file are silently skipped."],
            ["BR-06", "Scheduled runs process all rows in the Excel file regardless of date."],
            ["BR-07", "A processing log entry is always created, even when the run fails entirely."],
        ],
        col_widths=[1.2, 5.3]
    )

    # ── 7. Screen Inventory ───────────────────────────────────────────────────
    heading1(doc, "7.  Screen Inventory")
    divider(doc)
    add_styled_table(doc,
        ["Screen", "Navigation Label", "Key Actions"],
        [
            ["Dashboard",       "Dashboard",        "View daily stats, last run times, today's CRQs, recent logs, refresh"],
            ["CRQ List",        "CRQ List",          "Search by keyword, filter by status, view all CRQ records"],
            ["Ad-Hoc Approval", "Ad-Hoc Approval",   "Enter name and date range, trigger run, view result summary"],
            ["Processing Logs", "Processing Logs",   "View full run history with status and error messages"],
        ],
        col_widths=[1.8, 1.8, 3.0]
    )

    # ── 8. Non-Functional Requirements ────────────────────────────────────────
    heading1(doc, "8.  Non-Functional Requirements")
    divider(doc)
    add_styled_table(doc,
        ["NFR ID", "Category", "Description"],
        [
            ["NFR-01", "Performance",   "Each scheduled run must complete within the processing window before the next scheduled trigger."],
            ["NFR-02", "Performance",   "The dashboard must load within 3 seconds under normal operating conditions."],
            ["NFR-03", "Security",      "All external credentials (OneDrive, Remedy, SMTP) must be stored in environment configuration files, never in source code."],
            ["NFR-04", "Deployability", "The application must run as a single deployable JAR with no external web server required."],
            ["NFR-05", "Testability",   "A mock profile must allow full end-to-end demonstration without any external service credentials."],
            ["NFR-06", "Reliability",   "A single CRQ's Remedy API failure must not abort processing of the remaining CRQs."],
        ],
        col_widths=[1.0, 1.4, 4.1]
    )

    doc.save("docs/FSD.docx")
    print("FSD.docx created.")


# ─────────────────────────────────────────────────────────────────────────────
# TSD
# ─────────────────────────────────────────────────────────────────────────────

def build_tsd():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    cover_page(
        doc,
        "Technical Specification Document",
        "CRQ Approval Automation System",
        "1.0", "2026-05-11", "Draft"
    )

    # ── 1. Overview ───────────────────────────────────────────────────────────
    heading1(doc, "1.  System Overview")
    divider(doc)
    body(doc,
        "CRQ Approval is a monolithic full-stack application. The React single-page frontend is "
        "embedded inside the Spring Boot JAR at build time and served as static resources — "
        "no separate web server is required in production. All components run on a single "
        "port (8081).")
    doc.add_paragraph()
    body(doc, "High-level component interaction:", bold=True)
    for line in [
        "Browser (React SPA)  →  HTTP /api/crq/*",
        "Spring Boot (:8081)  →  H2 / PostgreSQL (database)",
        "Spring Boot          →  Microsoft Graph API (OneDrive Excel download)",
        "Spring Boot          →  Remedy REST API (CRQ status lookup)",
        "Spring Boot          →  SMTP Server (approval email)",
    ]:
        bullet(doc, line)

    # ── 2. Technology Stack ───────────────────────────────────────────────────
    heading1(doc, "2.  Technology Stack")
    divider(doc)

    heading2(doc, "2.1  Backend")
    add_styled_table(doc,
        ["Component", "Technology", "Version"],
        [
            ["Runtime",            "Java",                        "17"],
            ["Framework",          "Spring Boot",                 "3.2.5"],
            ["Build Tool",         "Apache Maven (wrapper)",      "3.x"],
            ["ORM",                "Spring Data JPA / Hibernate", "6.x"],
            ["Database (Dev)",     "H2 (file-based)",             "2.x"],
            ["Database (SIT)",     "PostgreSQL",                  "—"],
            ["Excel Parsing",      "Apache POI",                  "5.2.5"],
            ["OneDrive SDK",       "Microsoft Graph SDK",         "6.4.0"],
            ["Azure Auth",         "Azure Identity",              "1.12.2"],
            ["Code Generation",    "Lombok",                      "—"],
            ["Email",              "Spring Mail (Jakarta Mail)",  "—"],
            ["Scheduling",         "Spring @Scheduled",           "—"],
        ],
        col_widths=[2.0, 2.5, 2.0]
    )

    heading2(doc, "2.2  Frontend")
    add_styled_table(doc,
        ["Component", "Technology", "Version"],
        [
            ["Framework",         "React",            "18.3.1"],
            ["Build Tool",        "Vite",             "7.0.0"],
            ["HTTP Client",       "Axios",            "1.7.2"],
            ["Styling",           "Vanilla CSS",      "—"],
            ["State Management",  "React Hooks",      "—"],
            ["Routing",           "None (useState tab switching)", "—"],
        ],
        col_widths=[2.0, 2.5, 2.0]
    )

    # ── 3. Spring Profiles ────────────────────────────────────────────────────
    heading1(doc, "3.  Spring Profile Configuration")
    divider(doc)
    body(doc,
        "Three Spring profiles control which service implementations are active and which "
        "database is used. Switching profiles requires no code changes.")
    doc.add_paragraph()
    add_styled_table(doc,
        ["Profile", "Database", "OneDrive", "Remedy API", "Email", "Port"],
        [
            ["(default)", "H2 file ./data/crqdb",     "Real MS Graph API", "Real Remedy REST", "Real SMTP",    "8081"],
            ["mock",      "H2 in-memory (reset on restart)", "Generated Excel",  "Hardcoded statuses", "Console log only", "8081"],
            ["sit",       "PostgreSQL",                "Real MS Graph API", "Real Remedy REST", "Real SMTP",    "8081"],
        ],
        col_widths=[1.0, 1.8, 1.5, 1.4, 1.3, 0.6]
    )
    doc.add_paragraph()
    body(doc, "Activate mock profile:  ", bold=False)
    code_para(doc, "JAVA_HOME=$(/usr/libexec/java_home -v 17) ./mvnw spring-boot:run -Dspring-boot.run.profiles=mock")

    # ── 4. Package Structure ──────────────────────────────────────────────────
    heading1(doc, "4.  Backend Package Structure")
    divider(doc)
    add_styled_table(doc,
        ["Package / Class", "Role"],
        [
            ["CrqApprovalApplication",       "Entry point — @SpringBootApplication, @EnableScheduling"],
            ["controller/CrqController",     "REST endpoints at /api/crq (dashboard, list, logs, adhoc, run-now)"],
            ["service/CrqService",           "Core orchestration: download → parse → filter → check status → email → log"],
            ["service/OneDrivePort",         "Interface: download Excel InputStream from OneDrive"],
            ["service/OneDriveService",      "Real implementation using Microsoft Graph SDK (@Profile(\"!mock\"))"],
            ["service/RemedyPort",           "Interface: fetch CRQ status and check approval"],
            ["service/RemedyService",        "Real implementation calling Remedy REST API (@Profile(\"!mock\"))"],
            ["service/EmailPort",            "Interface: send approval notification email"],
            ["service/EmailService",         "Real implementation using Spring Mail (@Profile(\"!mock\"))"],
            ["service/ExcelService",         "Apache POI parser — reads rows, skips header, returns ExcelCrqRow list"],
            ["repository/CrqRepository",     "Spring Data JPA for Crq entity; custom queries for today/batch lookups"],
            ["repository/ProcessingLogRepository", "Spring Data JPA for ProcessingLog; top-10 and all-log queries"],
            ["model/Crq",                    "JPA entity: all CRQ fields + approval/email status + batch metadata"],
            ["model/ProcessingLog",          "JPA entity: run audit record with counts, status, error message"],
            ["dto/CrqDto",                   "API response DTO — static from(Crq) factory method"],
            ["dto/DashboardDto",             "Aggregated dashboard response — static from() factory method"],
            ["dto/ProcessingLogDto",         "Log response DTO — static from(ProcessingLog) factory method"],
            ["mock/MockDataInitializer",     "Seeds demo data on startup (@Profile(\"mock\")) via ApplicationRunner"],
            ["mock/MockOneDriveService",     "Generates 10-row in-memory Excel workbook"],
            ["mock/MockRemedyService",       "Returns hardcoded statuses — 6 of 10 CRQs are approved"],
            ["mock/MockEmailService",        "Logs email content to console instead of sending SMTP"],
            ["config/AppConfig",             "Spring bean definitions (RestTemplate, etc.)"],
            ["config/WebConfig",             "CORS configuration for /api/**"],
            ["scheduler/CrqScheduler",       "@Scheduled trigger — fires CrqService.runScheduledJob() at 5:30 PM"],
        ],
        col_widths=[2.5, 4.0]
    )

    # ── 5. External Service Ports ─────────────────────────────────────────────
    heading1(doc, "5.  External Service Abstraction (Port Pattern)")
    divider(doc)
    body(doc,
        "All three external integrations are hidden behind Java interfaces. CrqService depends "
        "only on the interfaces, so swapping mock and real implementations requires no changes "
        "to the core business logic.")
    doc.add_paragraph()

    heading2(doc, "5.1  OneDrivePort")
    add_styled_table(doc,
        ["Method", "Returns", "Description"],
        [
            ["downloadExcelFile()", "InputStream", "Downloads the Excel file from the configured OneDrive path"],
            ["getFileLastModifiedTime()", "OffsetDateTime", "Returns the last-modified timestamp of the file"],
        ],
        col_widths=[2.5, 1.3, 2.7]
    )
    body(doc, "Real implementation: authenticates using Azure ClientSecretCredential. Resolves the "
         "OneDrive drive ID first via the Graph SDK, then downloads the file by path.")

    heading2(doc, "5.2  RemedyPort")
    add_styled_table(doc,
        ["Method", "Returns", "Description"],
        [
            ["getCrqStatus(crqNumber)", "String", "Calls Remedy REST API with bearer token; returns current status string"],
            ["isApproved(status)", "boolean", "Returns true if status equals the configured approved-status value"],
        ],
        col_widths=[2.5, 1.0, 3.0]
    )

    heading2(doc, "5.3  EmailPort")
    add_styled_table(doc,
        ["Method", "Description"],
        [
            ["sendApprovalEmail(List<Crq>)", "Sends one consolidated email listing all approved CRQs"],
            ["sendApprovalEmail(Crq)", "Sends a single-CRQ notification (overload)"],
        ],
        col_widths=[2.5, 4.0]
    )

    # ── 6. Database Schema ────────────────────────────────────────────────────
    heading1(doc, "6.  Database Schema")
    divider(doc)

    heading2(doc, "6.1  Table: crq")
    add_styled_table(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id",                    "BIGINT",       "PK, AUTO_INCREMENT", "Surrogate primary key"],
            ["crq_number",            "VARCHAR",      "NOT NULL",           "CRQ identifier from Excel"],
            ["title",                 "VARCHAR",      "",                   "CRQ title"],
            ["assignee",              "VARCHAR",      "",                   "Assigned person or team"],
            ["description",           "VARCHAR",      "",                   "CRQ description"],
            ["remedy_status",         "VARCHAR",      "",                   "Status fetched from Remedy API"],
            ["approved",              "BOOLEAN",      "",                   "True if remedy_status == approved-status config"],
            ["email_sent",            "BOOLEAN",      "",                   "True after notification email is sent"],
            ["email_sent_at",         "TIMESTAMP",    "",                   "When the notification email was sent"],
            ["processed_at",          "TIMESTAMP",    "",                   "When this record was created"],
            ["last_updated_in_excel", "TIMESTAMP",    "",                   "Value from Excel column 4 (Last Updated)"],
            ["batch_type",            "VARCHAR",      "",                   "SCHEDULED or ADHOC"],
            ["batch_run_at",          "TIMESTAMP",    "",                   "Batch run identifier / start time"],
        ],
        col_widths=[1.8, 1.0, 1.4, 2.4]
    )

    heading2(doc, "6.2  Table: processing_log")
    add_styled_table(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id",              "BIGINT",       "PK, AUTO_INCREMENT", "Surrogate primary key"],
            ["run_at",          "TIMESTAMP",    "",                   "When the run started"],
            ["batch_type",      "VARCHAR",      "",                   "SCHEDULED or ADHOC"],
            ["total_crqs_read", "INT",          "",                   "Number of CRQ rows parsed from Excel"],
            ["approved_count",  "INT",          "",                   "Number of CRQs with approved status"],
            ["emails_sent",     "INT",          "",                   "1 if consolidated email was sent, else 0"],
            ["status",          "VARCHAR",      "",                   "SUCCESS, PARTIAL, or FAILED"],
            ["error_message",   "VARCHAR(2000)","",                   "Error detail if status is not SUCCESS"],
            ["triggered_by",    "VARCHAR",      "",                   "\"SCHEDULER\" or the ad-hoc user's name"],
        ],
        col_widths=[1.8, 1.2, 1.2, 2.4]
    )

    heading2(doc, "6.3  DDL Notes")
    add_styled_table(doc,
        ["Environment", "DDL Mode", "Notes"],
        [
            ["Development (H2 file)",    "update",      "Schema updated automatically; data persists across restarts"],
            ["Mock (H2 in-memory)",       "create-drop", "Schema recreated on every restart; demo data seeded by MockDataInitializer"],
            ["SIT (PostgreSQL)",          "update",      "Schema managed by Hibernate; no migration tool used"],
        ],
        col_widths=[2.0, 1.3, 3.2]
    )

    # ── 7. Processing Pipeline ────────────────────────────────────────────────
    heading1(doc, "7.  Processing Pipeline")
    divider(doc)
    body(doc,
        "Both scheduled and ad-hoc runs share a single pipeline method "
        "(CrqService.runJob). The only difference is that ad-hoc runs apply a "
        "date-time range filter at step 3.")
    doc.add_paragraph()

    steps = [
        ("Step 1 — Download",      "OneDrivePort.downloadExcelFile() returns an InputStream of the Excel file."),
        ("Step 2 — Parse",         "ExcelService.parseExcel() reads all rows using Apache POI, skipping the header row and any row with a blank CRQ number. Returns a list of ExcelCrqRow objects."),
        ("Step 3 — Filter (Ad-Hoc only)", "Rows are filtered to those where lastUpdatedInExcel falls within the caller-supplied fromDateTime–toDateTime range."),
        ("Step 4 — Status Check",  "For each row, RemedyPort.getCrqStatus(crqNumber) is called. If approved, the approved flag is set to true. A Crq entity is built and saved to the database."),
        ("Step 5 — Email",         "If any CRQs are approved, EmailPort.sendApprovalEmail(approvedList) sends one consolidated email. The emailSent flag and emailSentAt timestamp are then saved on each approved record."),
        ("Step 6 — Log",           "A ProcessingLog entity is created with counts (totalCrqsRead, approvedCount, emailsSent) and status (SUCCESS / PARTIAL / FAILED). Individual Remedy errors mark the run PARTIAL; a top-level exception marks it FAILED."),
    ]

    add_styled_table(doc,
        ["Step", "Description"],
        steps,
        col_widths=[2.0, 4.5]
    )

    # ── 8. REST API ────────────────────────────────────────────────────────────
    heading1(doc, "8.  REST API Specification")
    divider(doc)
    body(doc, "Base path: /api/crq")
    doc.add_paragraph()

    add_styled_table(doc,
        ["Method", "Path", "Request Body", "Response", "Description"],
        [
            ["GET",  "/dashboard", "—",
             "DashboardDto",
             "Today's stats, last run times, recent CRQs, recent logs (top 10)"],
            ["GET",  "/list", "—",
             "List<CrqDto>",
             "All CRQ records ordered by processedAt descending"],
            ["GET",  "/logs", "—",
             "List<ProcessingLogDto>",
             "All processing log entries ordered by runAt descending"],
            ["POST", "/adhoc",
             "{ triggeredBy, fromDateTime, toDateTime }",
             "ProcessingLogDto",
             "Triggers ad-hoc run for CRQs updated within the given date-time range"],
            ["POST", "/run-now", "—",
             "ProcessingLogDto",
             "Forces a full scheduled run immediately (admin/testing use)"],
        ],
        col_widths=[0.6, 1.1, 1.8, 1.4, 2.6]
    )

    doc.add_paragraph()
    heading2(doc, "8.1  POST /adhoc — Request Validation")
    add_styled_table(doc,
        ["Field", "Type", "Required", "Validation Rule"],
        [
            ["triggeredBy",  "String",    "Yes", "Must not be blank"],
            ["fromDateTime", "ISO-8601 string", "Yes", "Must be a valid local date-time"],
            ["toDateTime",   "ISO-8601 string", "Yes", "Must be a valid local date-time and after fromDateTime"],
        ],
        col_widths=[1.4, 1.5, 0.9, 2.8]
    )
    body(doc, "Returns HTTP 400 with a validation error message if any rule is violated.")

    # ── 9. Frontend Architecture ───────────────────────────────────────────────
    heading1(doc, "9.  Frontend Architecture")
    divider(doc)

    heading2(doc, "9.1  Component Hierarchy")
    add_styled_table(doc,
        ["Component", "File", "Responsibility"],
        [
            ["App",              "App.jsx",                        "Root shell: sidebar navigation, active-tab state, top bar"],
            ["Dashboard",        "components/Dashboard.jsx",       "Stats cards, last run info, today's CRQs table, recent logs table"],
            ["CrqTable",         "components/CrqTable.jsx",        "Full CRQ list with search box and status filter buttons"],
            ["AdHocApproval",    "components/AdHocApproval.jsx",   "Ad-hoc run form with validation and result display"],
            ["ProcessingLogs",   "components/ProcessingLogs.jsx",  "Full processing log history table"],
            ["API Layer",        "services/api.js",                "Axios instance wrapping all five backend API calls"],
        ],
        col_widths=[1.5, 2.3, 2.8]
    )

    heading2(doc, "9.2  State Management")
    body(doc,
        "The application uses React's built-in useState hook exclusively. "
        "App.jsx holds the activeTab state. Each component manages its own loading, "
        "data, and error states. No external state library (Redux, Zustand, or Context) is used.")

    heading2(doc, "9.3  API Layer")
    add_styled_table(doc,
        ["Function", "HTTP Call", "Returns"],
        [
            ["getDashboard()",                    "GET /dashboard",     "DashboardDto"],
            ["getAllCrqs()",                       "GET /list",          "CrqDto[]"],
            ["getLogs()",                          "GET /logs",          "ProcessingLogDto[]"],
            ["triggerAdhoc(by, from, to)",         "POST /adhoc",        "ProcessingLogDto"],
            ["runNow()",                           "POST /run-now",      "ProcessingLogDto"],
        ],
        col_widths=[2.5, 1.8, 2.2]
    )
    body(doc,
        "In development, the Vite dev server proxies all /api requests to http://localhost:8081. "
        "In production, the React build is served by Spring Boot on the same origin, so no proxy is needed.")

    heading2(doc, "9.4  Client-Side Filtering")
    body(doc,
        "CRQ list filtering and search are performed entirely in the browser after fetching the "
        "full dataset. No server-side pagination is implemented. This is adequate for the current "
        "data volume but will not scale to very large datasets.")

    # ── 10. Configuration Reference ───────────────────────────────────────────
    heading1(doc, "10.  Configuration Properties Reference")
    divider(doc)
    add_styled_table(doc,
        ["Property", "Description"],
        [
            ["server.port",                    "HTTP port — default 8081"],
            ["crq.scheduler.cron",             "Cron expression for scheduled run — default: 0 30 17 * * ? (5:30 PM daily)"],
            ["crq.remedy.base-url",            "Remedy REST API base URL"],
            ["crq.remedy.username",            "Remedy API username"],
            ["crq.remedy.password",            "Remedy API password"],
            ["crq.remedy.approved-status",     "Status string that means approved (e.g. \"Request in Change\")"],
            ["crq.onedrive.tenant-id",         "Azure Active Directory tenant ID"],
            ["crq.onedrive.client-id",         "Azure AD application (client) ID"],
            ["crq.onedrive.client-secret",     "Azure AD client secret"],
            ["crq.onedrive.user-email",        "Email of the OneDrive owner (used to locate the drive)"],
            ["crq.onedrive.file-path",         "Path to the Excel file within OneDrive"],
            ["crq.email.from",                 "Sender email address"],
            ["crq.email.to",                   "Recipient email address(es)"],
            ["crq.email.subject",              "Email subject line"],
            ["spring.mail.host",               "SMTP server hostname"],
            ["spring.mail.port",               "SMTP server port"],
            ["spring.mail.username",           "SMTP authentication username"],
            ["spring.mail.password",           "SMTP authentication password"],
            ["cors.allowed-origins",           "Comma-separated list of allowed CORS origins"],
            ["spring.datasource.url",          "(SIT) PostgreSQL JDBC URL"],
            ["spring.datasource.username",     "(SIT) PostgreSQL username"],
            ["spring.datasource.password",     "(SIT) PostgreSQL password"],
        ],
        col_widths=[3.0, 3.5]
    )

    # ── 11. Build and Deployment ───────────────────────────────────────────────
    heading1(doc, "11.  Build and Deployment")
    divider(doc)

    heading2(doc, "11.1  Development Commands")
    add_styled_table(doc,
        ["Command", "Purpose"],
        [
            ["JAVA_HOME=$(...) ./mvnw spring-boot:run -Dspring-boot.run.profiles=mock",
             "Run backend in mock mode (no external credentials needed)"],
            ["JAVA_HOME=$(...) ./mvnw spring-boot:run",
             "Run backend with real credentials from application.properties"],
            ["JAVA_HOME=$(...) ./mvnw compile",
             "Compile only (verify build without running)"],
            ["cd frontend && npm install && npm run dev",
             "Start React dev server on port 5173 (proxies /api to :8081)"],
            ["cd frontend && npm run build",
             "Build React app into frontend/dist/"],
            ["bash build-sit.sh",
             "Build full SIT deployment package — crq-approval-sit-1.0.0.zip"],
        ],
        col_widths=[3.2, 3.3]
    )

    heading2(doc, "11.2  SIT Deployment Package")
    body(doc, "build-sit.sh produces crq-approval-sit-1.0.0.zip containing:")
    add_styled_table(doc,
        ["File", "Description"],
        [
            ["crq-approval-1.0.0.jar",    "Fat JAR — Spring Boot + React embedded as static resources"],
            ["application-sit.properties","Environment configuration template (fill in credentials before deploy)"],
            ["start.sh",                  "Starts the JAR with spring.profiles.active=sit"],
            ["stop.sh",                   "Stops the running process"],
        ],
        col_widths=[2.5, 4.0]
    )
    body(doc,
        "The Maven package phase copies frontend/dist/ into target/classes/static/ via the "
        "maven-resources-plugin. Spring Boot's built-in static resource handler then serves the "
        "React app at the root path /.")

    # ── 12. Known Technical Constraints ───────────────────────────────────────
    heading1(doc, "12.  Known Technical Constraints")
    divider(doc)
    add_styled_table(doc,
        ["Constraint", "Detail"],
        [
            ["Java 17 required",
             "The system's default Maven uses Java 8. Every mvnw command must be prefixed with JAVA_HOME=$(/usr/libexec/java_home -v 17)."],
            ["H2 2.x URL restriction",
             "Cannot combine AUTO_SERVER=TRUE with DB_CLOSE_ON_EXIT=FALSE in the H2 JDBC URL — H2 2.x rejects this combination."],
            ["Microsoft Graph SDK v6",
             "users().byUserId().drive() returns a limited builder with no root() or items(). Drive ID must be resolved first via .drive().get(), then use drives().byDriveId(id).items().byDriveItemId(\"root:/path:\")."],
            ["Hibernate 6 JPQL",
             "The DATE() function in JPQL returns Object, not LocalDate, breaking type-safe comparisons. Use findByProcessedAtBetween(LocalDateTime, LocalDateTime) instead."],
            ["No test suite",
             "No JUnit or Vitest tests exist in the repository. mvnw test and npm test have nothing to run."],
            ["No linting tools",
             "No Checkstyle, SpotBugs, ESLint, or Prettier is configured."],
            ["Client-side filtering",
             "All CRQ records are loaded into the browser on the CRQ List screen. Filtering is in-memory. This will not scale well as record volume grows."],
            ["react-router-dom",
             "The react-router-dom package is installed as a dependency but is not used. All navigation is implemented as useState tab switching in App.jsx."],
        ],
        col_widths=[1.9, 4.6]
    )

    doc.save("docs/TSD.docx")
    print("TSD.docx created.")


if __name__ == "__main__":
    build_fsd()
    build_tsd()
